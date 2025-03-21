import requests
from docx import Document
import winshell
from pathlib import Path



def export_to_word(routes_data, total_cost, filename):
    # Obtém o caminho real da área de trabalho usando `winshell`
    desktop_path = Path(winshell.desktop()) / filename

    doc = Document()
    doc.add_heading("Resumo das Rotas", level=1)

    for i, route in enumerate(routes_data, start=1):
        doc.add_heading(f"Rota {i}", level=2)
        doc.add_paragraph(f"Percurso → {' → '.join(route['optimized_route'])}")
        doc.add_paragraph(f"Distância: {route['distance']:.2f} km")
        doc.add_paragraph(f"Tempo estimado: {route['duration']:.2f} min")
        doc.add_paragraph(f"Custo do combustível: R$ {route['fuel_cost']:.2f}")
        doc.add_paragraph(f"Pedágios: R$ {route['toll_cost']:.2f}")
        doc.add_paragraph(f"Custo Total: R$ {route['total_cost']:.2f}")
        doc.add_paragraph("")

    doc.add_heading("Custo Total de Todas as Rotas", level=2)
    doc.add_paragraph(f"R$ {total_cost:.2f}")

    try:
        doc.save(desktop_path)
        print(f"Arquivo salvo com sucesso: {desktop_path}")
    except PermissionError:
        print("Erro: O arquivo já está aberto no Word. Feche o arquivo e tente novamente.")
    except FileNotFoundError:
        print(f"Erro: Caminho inválido {desktop_path}. Verifique se a área de trabalho existe.")
    except Exception as e:
        print(f"Erro inesperado ao salvar o arquivo: {e}")