import requests
from docx import Document
import winshell
from pathlib import Path
from kivy.graphics import Color, RoundedRectangle

# Configuração das APIs
TOLLGURU_API_KEY = "rfGDRF7rNhMP3nDt9mJhgM6f9N6rb8d6"
GOOGLE_MAPS_API_KEY = "AIzaSyDaCAYGqy0ZHkh4FddsYajv1V7b50mIlng"


# valores fixos
FUEL_PRICE = 6.20  # R$/L
FUEL_EFFICIENCY = 12  # km/L
    
def __init__(self, **kwargs):
        super().__init__(**kwargs)
        self.waypoint_limit = 23  # Padrão: Listagem Padrão (23 waypoints)

def set_bate_volta(self, instance):
    self.waypoint_limit = 1
    print("Modo Bate e Volta ativado (1 waypoint por rota).")

    # Atualiza a cor do botão bate e volta
    self.update_button_color(self.ids.btn_bate_volta, (0, 1, 0, 1))  # Verde
    self.update_button_color(self.ids.btn_padrao, (0.13, 0.59, 0.95, 1))  # Azul (padrão)

def set_padrao(self, instance):
    self.waypoint_limit = 23
    print("Modo Padrão ativado (23 waypoints por rota).")

    # Atualiza a cor do botão padrão
    self.update_button_color(self.ids.btn_bate_volta, (0.13, 0.59, 0.95, 1))  # Azul (padrão)
    self.update_button_color(self.ids.btn_padrao, (0, 1, 0, 1))  # Verde

def update_button_color(self, button, color):
    """ Atualiza a cor de fundo do botão com Canvas e RoundedRectangle """
    button.canvas.before.clear()  # Limpa o desenho anterior
    with button.canvas.before:
        Color(*color)  # Define a nova cor
        RoundedRectangle(pos=button.pos, size=button.size, radius=[25])  # Mantém as bordas arredondadas
            
def start_calculation(self):
        origin = self.ids.origin.text
        destination = self.ids.destination.text
        waypoints = [w.strip() for w in self.ids.waypoints.text.split('/') if w.strip()]
        
        filename = self.ids.servico.text.strip() or "Rotas_Calculadas"
        filename = "".join(c if c.isalnum() or c in (" ", ".", "_") else "_" for c in filename) + ".docx"

        self.calculate_routes(origin, destination, waypoints, filename)

def split_waypoints(self, waypoints):
    return [waypoints[i:i + self.waypoint_limit] for i in range(0, len(waypoints), self.waypoint_limit)]

def get_route(self, origin, destination, waypoints):
    url = f"https://maps.googleapis.com/maps/api/directions/json?origin={origin}&destination={destination}&waypoints=optimize:true|{'|'.join(waypoints)}&key={GOOGLE_MAPS_API_KEY}"
    
    try:
        response = requests.get(url, timeout=10)
        response.raise_for_status()
        data = response.json()

        if data.get("status") == "OK":
            route = data["routes"][0]["legs"]
            distance = sum(leg["distance"]["value"] for leg in route) / 1000
            duration = sum(leg["duration"]["value"] for leg in route) / 60
            optimized_order = [origin] + [waypoints[i] for i in data["routes"][0]["waypoint_order"]] + [destination]             
            return distance, duration, optimized_order
    except requests.exceptions.RequestException as e:
        print(f"Erro na requisição: {e}")
        
    return None, None, None

def get_toll_cost(self, origin, destination, waypoints):
    url = "https://apis.tollguru.com/toll/v2/origin-destination-waypoints"
    headers = {"x-api-key": TOLLGURU_API_KEY, "Content-Type": "application/json"}
    payload = {
        "from": {"address": origin},
        "to": {"address": destination},
        "serviceProvider": "here",
        "waypoints": [{"address": w} for w in waypoints],
        "vehicle": {"type": "2AxlesAuto"},
        "departure_time": "now",
        "fuelType": "gasoline",
        "fuelEfficiency": {"value": FUEL_EFFICIENCY, "unit": "kmpl"},
        "currency": "BRL"
    }
        
    response = requests.post(url, json=payload, headers=headers)
        
    if response.status_code == 200:
        data = response.json()
        if "routes" in data and data["routes"]:
            costs = data["routes"][0].get("costs", {})
            return max(filter(None, [costs.get("tag", 0), costs.get("cash", 0), costs.get("licensePlate", 0)]), default=0)
    return 0

def calculate_routes(self, origin, destination, waypoints, filename):
    total_cost = 0
    routes_data = []
    # Supondo que 'self' seja a tela que contém a função split_waypoints:
    waypoints_groups = self.split_waypoints(waypoints)
    # Verifica se waypoints_groups é realmente uma lista
    if waypoints_groups is None:
        waypoints_groups = []  # Ou lance uma mensagem de erro informativa

    for group in waypoints_groups:
        # Lógica para calcular cada rota...
        distance, duration, optimized_route = self.get_route(origin, destination, group)
        toll_cost = self.get_toll_cost(origin, destination, group)
        if distance:
            fuel_needed = distance / FUEL_EFFICIENCY
            fuel_cost = fuel_needed * FUEL_PRICE
            route_total = fuel_cost + toll_cost
            total_cost += route_total
            routes_data.append({
                "optimized_route": optimized_route,
                "distance": distance,
                "duration": duration,
                "fuel_cost": fuel_cost,
                "toll_cost": toll_cost,
                "total_cost": route_total
            })
    export_to_word(routes_data, total_cost, filename)

def export_to_word(routes_data, total_cost, filename):
    # Obtém o caminho real da área de trabalho usando `winshell`
    desktop_path = Path(winshell.desktop()) / filename

    doc = Document()
    doc.add_heading("Resumo das Rotas", level=1)

    for i, route in enumerate(routes_data, start=1):
        doc.add_heading(f"Rota {i}", level=2)
        doc.add_paragraph(f"Percurso → {' → '.join(route['optimized_route'])}")
        doc.add_paragraph(f"Distância: {route['distance']:.2f} km")
        doc.add_paragraph(f"Tempo estimado: {route['duration']:.2f} min")
        doc.add_paragraph(f"Custo do combustível: R$ {route['fuel_cost']:.2f}")
        doc.add_paragraph(f"Pedágios: R$ {route['toll_cost']:.2f}")
        doc.add_paragraph(f"Custo Total: R$ {route['total_cost']:.2f}")
        doc.add_paragraph("")

    doc.add_heading("Custo Total de Todas as Rotas", level=2)
    doc.add_paragraph(f"R$ {total_cost:.2f}")

    try:
        doc.save(desktop_path)
        print(f"Arquivo salvo com sucesso: {desktop_path}")
    except PermissionError:
        print("Erro: O arquivo já está aberto no Word. Feche o arquivo e tente novamente.")
    except FileNotFoundError:
        print(f"Erro: Caminho inválido {desktop_path}. Verifique se a área de trabalho existe.")
    except Exception as e:
        print(f"Erro inesperado ao salvar o arquivo: {e}")
 

def all_reset(self):
    self.ids.servico.text = ""
    self.ids.origin.text = ""
    self.ids.destination.text = ""
    self.ids.waypoints.text = ""